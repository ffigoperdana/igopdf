#!/usr/bin/env python3
import argparse
import io
import json
import os
import shutil
import subprocess
import sys
import tempfile

import fitz
import pytesseract
from docx import Document
from docx.shared import Inches
from pdf2docx import Converter
from PIL import Image


def emit(payload):
    print(json.dumps(payload), flush=True)


def fail(code):
    emit({"type": "error", "code": code})
    raise RuntimeError(code)


def normalize_with_qpdf(source, directory):
    target = os.path.join(directory, "qpdf-repaired.pdf")
    result = subprocess.run(
        ["qpdf", "--warning-exit-0", "--object-streams=generate", source, target],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        check=False,
    )
    return target if result.returncode == 0 and os.path.exists(target) else source


def normalize_with_ghostscript(source, directory):
    target = os.path.join(directory, "ghostscript-normalized.pdf")
    result = subprocess.run(
        [
            "gs", "-dSAFER", "-dBATCH", "-dNOPAUSE", "-sDEVICE=pdfwrite",
            "-dCompatibilityLevel=1.6", "-dPDFSETTINGS=/prepress",
            "-sOutputFile=" + target, source,
        ],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        check=False,
    )
    return target if result.returncode == 0 and os.path.exists(target) else None


def convert_editable(source, output, workspace):
    emit({"type": "progress", "stage": "repairing", "progress": 12})
    repaired = normalize_with_qpdf(source, workspace)
    emit({"type": "progress", "stage": "converting", "progress": 32})
    try:
        converter = Converter(repaired)
        try:
            converter.convert(output)
        finally:
            converter.close()
        return
    except Exception:
        # Invalid embedded-font/xref references are common in PDFs exported by
        # third-party office systems. Rewriting the visual PDF removes those
        # broken objects, then pdf2docx gets one clean retry.
        emit({"type": "progress", "stage": "repairing", "progress": 50})
        normalized = normalize_with_ghostscript(repaired, workspace)
        if not normalized:
            fail("INVALID_PDF_STRUCTURE")
        try:
            converter = Converter(normalized)
            try:
                converter.convert(output)
            finally:
                converter.close()
        except Exception:
            fail("FONT_OR_LAYOUT_UNSUPPORTED")


def page_size_inches(page):
    return page.rect.width / 72.0, page.rect.height / 72.0


def convert_visual(document, output):
    docx = Document()
    total = len(document)
    for index, page in enumerate(document):
        width, height = page_size_inches(page)
        section = docx.sections[0] if index == 0 else docx.add_section()
        section.page_width = Inches(width)
        section.page_height = Inches(height)
        section.left_margin = Inches(0.18)
        section.right_margin = Inches(0.18)
        section.top_margin = Inches(0.18)
        section.bottom_margin = Inches(0.18)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        image = io.BytesIO(pix.tobytes("png"))
        docx.add_picture(image, width=Inches(max(0.1, width - 0.36)))
        emit({"type": "progress", "stage": "rendering", "progress": 18 + 75 * (index + 1) / total, "currentPage": index + 1, "totalPages": total})
    docx.save(output)


def convert_ocr(document, output):
    docx = Document()
    total = len(document)
    for index, page in enumerate(document):
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(image, lang="ind+eng")
        if text.strip():
            for paragraph in text.splitlines():
                docx.add_paragraph(paragraph)
        else:
            docx.add_paragraph("")
        if index < total - 1:
            docx.add_page_break()
        emit({"type": "progress", "stage": "ocr", "progress": 18 + 75 * (index + 1) / total, "currentPage": index + 1, "totalPages": total})
    docx.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--mode", choices=["editable", "ocr", "visual"], required=True)
    parser.add_argument("--max-pages", type=int, required=True)
    args = parser.parse_args()
    workspace = tempfile.mkdtemp(prefix="igo-docx-")
    try:
        emit({"type": "progress", "stage": "validating", "progress": 5})
        try:
            document = fitz.open(args.input)
        except Exception:
            fail("INVALID_PDF_STRUCTURE")
        try:
            if document.needs_pass:
                fail("ENCRYPTED_PDF")
            total = len(document)
            if total <= 0:
                fail("INVALID_PDF_STRUCTURE")
            if total > args.max_pages:
                fail("PAGE_LIMIT")
            emit({"type": "progress", "stage": "analyzing", "progress": 10, "currentPage": 0, "totalPages": total})
            if args.mode == "editable":
                document.close()
                convert_editable(args.input, args.output, workspace)
            elif args.mode == "ocr":
                convert_ocr(document, args.output)
            else:
                convert_visual(document, args.output)
            emit({"type": "progress", "stage": "packaging", "progress": 95, "currentPage": total, "totalPages": total})
        finally:
            if not document.is_closed:
                document.close()
    except RuntimeError:
        sys.exit(1)
    except MemoryError:
        emit({"type": "error", "code": "MEMORY_LIMIT"})
        sys.exit(1)
    except Exception:
        emit({"type": "error", "code": "PROCESS_FAILED"})
        sys.exit(1)
    finally:
        shutil.rmtree(workspace, ignore_errors=True)


if __name__ == "__main__":
    main()
